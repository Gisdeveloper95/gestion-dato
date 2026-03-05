#!/usr/bin/env python3
"""
Módulo de Consultas IA - Sistema Inteligente
Analiza archivos, busca por keywords, genera respuestas con contexto
"""

import os
import sys
import sqlite3
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import requests
from datetime import datetime

# Importar utilidades
sys.path.insert(0, str(Path(__file__).parent.parent))
from utils.common import Config, DatabaseManager


class AIQuerySystem:
    """Sistema de consultas inteligentes con IA"""

    def __init__(self, config_path: Optional[str] = None):
        self.config = Config(config_path)
        self.db = DatabaseManager(self.config)
        self.municipios_db_path = Path(__file__).parent / "municipios.db"

        # API Keys (agregar a config.json)
        self.groq_api_key = self.config.get('ai', 'groq_api_key', default='')
        self.groq_url = "https://api.groq.com/openai/v1/chat/completions"

    # ==================== BÚSQUEDA DE MUNICIPIOS ====================

    def buscar_municipio(self, nombre: str) -> Optional[Dict]:
        """Busca un municipio por nombre o código DANE y retorna su información"""
        if not self.municipios_db_path.exists():
            return None

        try:
            import unicodedata

            def normalizar(texto):
                """Elimina acentos y convierte a mayúsculas"""
                nfkd = unicodedata.normalize('NFKD', texto)
                return ''.join([c for c in nfkd if not unicodedata.combining(c)]).upper()

            conn = sqlite3.connect(self.municipios_db_path)
            cursor = conn.cursor()

            # Verificar si es un código numérico
            if nombre.strip().isdigit():
                # Buscar por código DANE exacto
                cursor.execute("SELECT CODIGO_DANE, MUNICIPIO, DEPARTAMENTO FROM municipios WHERE CODIGO_DANE = ?", (nombre.strip(),))
                row = cursor.fetchone()
                conn.close()

                if row:
                    return {
                        'cod_dane': row[0],
                        'nombre': row[1],
                        'cod_depto': row[0][:2],
                        'nombre_depto': row[2]
                    }
                return None

            # Normalizar nombre de búsqueda
            nombre_norm = normalizar(nombre)

            # Buscar por nombre (sin acentos)
            cursor.execute("SELECT CODIGO_DANE, MUNICIPIO, DEPARTAMENTO FROM municipios")
            todos = cursor.fetchall()

            results = []
            for row in todos:
                mun_norm = normalizar(row[1])
                if nombre_norm in mun_norm:
                    results.append(row)
                    if len(results) >= 5:
                        break

            conn.close()

            if not results:
                return None

            # Si hay múltiples, retornar lista
            municipios = []
            for row in results:
                municipios.append({
                    'cod_dane': row[0],
                    'nombre': row[1],
                    'cod_depto': row[0][:2],
                    'nombre_depto': row[2]
                })

            return municipios[0] if len(municipios) == 1 else {'multiples': municipios}

        except Exception as e:
            print(f"Error buscando municipio: {e}")
            return None

    def extraer_municipios_de_texto(self, texto: str) -> List[Dict]:
        """Extrae nombres de municipios del texto usando la BD"""
        if not self.municipios_db_path.exists():
            return []

        try:
            conn = sqlite3.connect(self.municipios_db_path)
            cursor = conn.cursor()

            # Obtener todos los nombres de municipios
            cursor.execute("SELECT CODIGO_DANE, MUNICIPIO, SUBSTR(CODIGO_DANE, 1, 2), DEPARTAMENTO FROM municipios")
            todos_municipios = cursor.fetchall()
            conn.close()

            # Buscar en el texto
            encontrados = []
            texto_lower = texto.lower()

            for cod_dane, nombre, cod_depto, nombre_depto in todos_municipios:
                if nombre.lower() in texto_lower:
                    encontrados.append({
                        'cod_dane': cod_dane,
                        'nombre': nombre,
                        'cod_depto': cod_depto,
                        'nombre_depto': nombre_depto
                    })

            return encontrados

        except Exception as e:
            print(f"Error extrayendo municipios: {e}")
            return []

    # ==================== BÚSQUEDA INTELIGENTE DE ARCHIVOS ====================

    def buscar_archivos_por_keywords(self, keywords: List[str], cod_mpio: Optional[str] = None) -> List[Dict]:
        """
        Busca archivos por palabras clave en sus rutas
        Triangula usando múltiples keywords para mejorar precisión
        """
        try:
            # Construir query dinámica
            conditions = []
            params = []

            # Búsqueda por keywords en path
            for keyword in keywords:
                conditions.append("path_file ILIKE %s")
                params.append(f"%{keyword}%")

            # Filtrar por municipio si se especifica
            if cod_mpio:
                conditions.append("cod_mpio = %s")
                params.append(cod_mpio)

            where_clause = " AND ".join(conditions) if conditions else "1=1"

            # Buscar en PREOPERACION
            query_preo = f"""
                SELECT
                    'PREOPERACION' as origen,
                    path_file,
                    peso_memoria,
                    fecha_disposicion,
                    SUBSTR(path_file, 30, 5) as cod_mpio,
                    usuario_windows
                FROM lista_archivos_preo
                WHERE {where_clause}
                LIMIT 20
            """

            # Buscar en POSTOPERACION (simplificado)
            query_post = None  # Por ahora solo PREOPERACION

            resultados = []

            # Ejecutar búsqueda en PREO
            conn = self.db.get_connection()
            cursor = conn.cursor()
            cursor.execute(query_preo, tuple(params))
            for row in cursor.fetchall():
                peso_str = row[2] if row[2] else "0"
                # Convertir peso_memoria a MB (puede venir como "2.5 GB", "150 MB", etc)
                try:
                    if "GB" in peso_str.upper():
                        size_mb = float(peso_str.split()[0]) * 1024
                    elif "MB" in peso_str.upper():
                        size_mb = float(peso_str.split()[0])
                    else:
                        size_mb = 0
                except:
                    size_mb = 0

                resultados.append({
                    'origen': row[0],
                    'path': row[1],
                    'size_mb': size_mb,
                    'fecha': row[3].strftime('%Y-%m-%d') if row[3] else 'N/A',
                    'cod_mpio': row[4],
                    'propietario': row[5]
                })

            conn.close()

            return resultados

        except Exception as e:
            print(f"Error buscando archivos: {e}")
            return []

    def obtener_estadisticas_municipio(self, cod_mpio: str) -> Dict:
        """Obtiene estadísticas completas de un municipio"""
        try:
            # Búsqueda simplificada por path
            query = """
                SELECT
                    nombre_insumo,
                    COUNT(*) as total_archivos,
                    MAX(fecha_disposicion) as ultima_modificacion
                FROM lista_archivos_preo
                WHERE path_file LIKE %s
                GROUP BY nombre_insumo
                ORDER BY total_archivos DESC
            """

            conn = self.db.get_connection()
            cursor = conn.cursor()
            cursor.execute(query, (f"%{cod_mpio}%",))

            estadisticas = []
            total_archivos = 0

            for row in cursor.fetchall():
                categoria = row[0]
                archivos = row[1]
                fecha = row[2]

                estadisticas.append({
                    'categoria': categoria,
                    'archivos': archivos,
                    'peso_mb': 0,  # No disponible en este esquema
                    'ultima_modificacion': fecha.strftime('%Y-%m-%d') if fecha else 'N/A'
                })

                total_archivos += archivos

            conn.close()

            return {
                'cod_mpio': cod_mpio,
                'total_archivos': total_archivos,
                'peso_total_gb': 0,  # Calcular manualmente si se necesita
                'categorias': estadisticas
            }

        except Exception as e:
            print(f"Error obteniendo estadísticas: {e}")
            return {}

    # ==================== ANÁLISIS DE CONTENIDO ====================

    def analizar_archivo(self, ruta_archivo: str) -> Dict:
        """
        Analiza el contenido de un archivo según su tipo
        Soporta: Excel, CSV, PDF, Word, TXT
        """
        archivo = Path(ruta_archivo)

        if not archivo.exists():
            return {'error': 'Archivo no encontrado'}

        extension = archivo.suffix.lower()
        resultado = {
            'path': str(archivo),
            'nombre': archivo.name,
            'extension': extension,
            'size_mb': round(archivo.stat().st_size / (1024 * 1024), 2),
            'analisis': {}
        }

        try:
            if extension in ['.xlsx', '.xls']:
                resultado['analisis'] = self._analizar_excel(archivo)
            elif extension == '.csv':
                resultado['analisis'] = self._analizar_csv(archivo)
            elif extension == '.pdf':
                resultado['analisis'] = self._analizar_pdf(archivo)
            elif extension in ['.docx', '.doc']:
                resultado['analisis'] = self._analizar_word(archivo)
            elif extension in ['.txt', '.log']:
                resultado['analisis'] = self._analizar_texto(archivo)
            elif extension == '.shp':
                resultado['analisis'] = {'tipo': 'Shapefile (archivo GIS)', 'requiere_herramientas_especiales': True}
            else:
                resultado['analisis'] = {'tipo': 'Archivo binario', 'no_analizable': True}

        except Exception as e:
            resultado['analisis'] = {'error': str(e)}

        return resultado

    def _analizar_excel(self, archivo: Path) -> Dict:
        """Analiza archivo Excel"""
        try:
            import pandas as pd
            df = pd.read_excel(archivo, nrows=100)  # Solo primeras 100 filas

            return {
                'tipo': 'Excel',
                'filas': len(df),
                'columnas': len(df.columns),
                'nombres_columnas': list(df.columns),
                'muestra': df.head(5).to_dict('records')
            }
        except Exception as e:
            return {'error': f"No se pudo leer Excel: {e}"}

    def _analizar_csv(self, archivo: Path) -> Dict:
        """Analiza archivo CSV"""
        try:
            import pandas as pd
            df = pd.read_csv(archivo, nrows=100)

            return {
                'tipo': 'CSV',
                'filas': len(df),
                'columnas': len(df.columns),
                'nombres_columnas': list(df.columns),
                'muestra': df.head(5).to_dict('records')
            }
        except Exception as e:
            return {'error': f"No se pudo leer CSV: {e}"}

    def _analizar_pdf(self, archivo: Path) -> Dict:
        """Analiza archivo PDF"""
        try:
            import PyPDF2
            with open(archivo, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                texto_muestra = pdf.pages[0].extract_text()[:500]

                return {
                    'tipo': 'PDF',
                    'paginas': len(pdf.pages),
                    'muestra_texto': texto_muestra
                }
        except Exception as e:
            return {'error': f"No se pudo leer PDF: {e}"}

    def _analizar_word(self, archivo: Path) -> Dict:
        """Analiza archivo Word"""
        try:
            import docx
            doc = docx.Document(archivo)
            texto = '\n'.join([p.text for p in doc.paragraphs[:10]])

            return {
                'tipo': 'Word',
                'parrafos': len(doc.paragraphs),
                'muestra_texto': texto[:500]
            }
        except Exception as e:
            return {'error': f"No se pudo leer Word: {e}"}

    def _analizar_texto(self, archivo: Path) -> Dict:
        """Analiza archivo de texto"""
        try:
            with open(archivo, 'r', encoding='utf-8') as f:
                lineas = f.readlines()
                muestra = ''.join(lineas[:20])

            return {
                'tipo': 'Texto plano',
                'lineas': len(lineas),
                'muestra': muestra[:500]
            }
        except Exception as e:
            return {'error': f"No se pudo leer texto: {e}"}

    # ==================== CONSULTAS CON IA ====================

    def consultar_con_ia(self, pregunta: str, contexto: Optional[Dict] = None) -> str:
        """
        Envía pregunta a IA con contexto de BD
        Usa Groq (gratuito) para generar respuesta
        """
        if not self.groq_api_key:
            return "⚠️ No hay API key de Groq configurada. Agrega 'groq_api_key' en config.json"

        # Construir contexto
        system_prompt = """Eres un asistente experto en análisis de datos territoriales de Colombia.
Respondes preguntas sobre archivos, municipios, y datos geográficos.
Tus respuestas son claras, concisas y profesionales.
Cuando des estadísticas, usa formato amigable con emojis."""

        user_message = pregunta

        if contexto:
            user_message += f"\n\nCONTEXTO:\n{json.dumps(contexto, indent=2, ensure_ascii=False)}"

        try:
            # Leer configuración del modelo desde config.json
            model = self.config.get('ai', 'model', default='llama-3.3-70b-versatile')
            temperature = self.config.get('ai', 'temperature', default=0.7)
            max_tokens = self.config.get('ai', 'max_tokens', default=1000)

            response = requests.post(
                self.groq_url,
                headers={
                    "Authorization": f"Bearer {self.groq_api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_message}
                    ],
                    "temperature": temperature,
                    "max_tokens": max_tokens
                },
                timeout=30
            )

            if response.status_code == 200:
                data = response.json()
                return data['choices'][0]['message']['content']
            else:
                return f"❌ Error en API: {response.status_code} - {response.text}"

        except Exception as e:
            return f"❌ Error consultando IA: {str(e)}"

    def procesar_consulta_completa(self, pregunta: str) -> Dict:
        """
        Procesa una consulta completa:
        1. Extrae municipios mencionados
        2. Identifica keywords para búsqueda
        3. Busca archivos relevantes
        4. Genera respuesta con IA
        """
        resultado = {
            'pregunta': pregunta,
            'municipios': [],
            'archivos_encontrados': [],
            'respuesta_ia': '',
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }

        # 1. Extraer municipios
        municipios = self.extraer_municipios_de_texto(pregunta)
        resultado['municipios'] = municipios

        # 2. Extraer keywords (palabras relevantes)
        keywords = self._extraer_keywords(pregunta)

        # 3. Buscar archivos
        cod_mpio = municipios[0]['cod_dane'] if municipios else None
        archivos = self.buscar_archivos_por_keywords(keywords, cod_mpio)
        resultado['archivos_encontrados'] = archivos[:5]  # Máximo 5

        # 4. Construir contexto para IA
        contexto = {
            'municipios_mencionados': municipios,
            'archivos_relevantes': archivos[:5],
            'total_archivos_encontrados': len(archivos)
        }

        # Si hay municipio, agregar estadísticas
        if cod_mpio:
            stats = self.obtener_estadisticas_municipio(cod_mpio)
            contexto['estadisticas_municipio'] = stats

        # 5. Consultar IA
        resultado['respuesta_ia'] = self.consultar_con_ia(pregunta, contexto)

        return resultado

    def _extraer_keywords(self, texto: str) -> List[str]:
        """Extrae palabras clave relevantes del texto"""
        # Palabras irrelevantes
        stopwords = {
            'el', 'la', 'de', 'en', 'y', 'a', 'los', 'las', 'un', 'una',
            'por', 'para', 'con', 'del', 'al', 'que', 'es', 'tiene', 'hay'
        }

        # Limpiar y dividir
        palabras = texto.lower().split()
        keywords = [p.strip('.,?¿!¡()[]{}') for p in palabras if len(p) > 3 and p not in stopwords]

        return keywords[:5]  # Máximo 5 keywords


if __name__ == '__main__':
    # Test básico
    ai = AIQuerySystem()

    # Test búsqueda de municipio
    municipio = ai.buscar_municipio("Ibagué")
    print("Municipio:", municipio)

    # Test búsqueda de archivos
    archivos = ai.buscar_archivos_por_keywords(["shp", "cartografia"], "73001")
    print(f"Archivos encontrados: {len(archivos)}")
